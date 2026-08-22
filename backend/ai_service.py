import ast
import json
import operator
from collections.abc import Generator
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any
from zoneinfo import ZoneInfo, ZoneInfoNotFoundError

from openai import OpenAI

from . import database
from .config import settings
from .space_engine import SpaceRunMode, compile_space_system_prompt
from .workspaces import DEFAULT_WORKSPACE, WORKSPACES, validate_workspace


client = OpenAI(api_key=settings.openai_api_key)
MAX_TOOL_ROUNDS = 3
RAG_MIN_SIMILARITY = 0.30
MAX_EXTRACTED_CHARACTERS = 500_000
MAX_DOCUMENT_CHUNKS = 600

BASE_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "calculate",
            "description": "Evaluate a basic arithmetic expression safely.",
            "strict": True,
            "parameters": {
                "type": "object",
                "properties": {
                    "expression": {
                        "type": "string",
                        "description": "Arithmetic expression using numbers and +, -, *, /, //, %, **, parentheses.",
                    }
                },
                "required": ["expression"],
                "additionalProperties": False,
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_current_time",
            "description": "Get the current date and time in an IANA timezone.",
            "strict": True,
            "parameters": {
                "type": "object",
                "properties": {
                    "timezone": {
                        "type": "string",
                        "description": "IANA timezone such as Australia/Sydney or Asia/Shanghai.",
                    }
                },
                "required": ["timezone"],
                "additionalProperties": False,
            },
        },
    },
]

FINANCE_TOOL = {
    "type": "function",
    "function": {
        "name": "calculate_financial_metric",
        "description": (
            "Calculate a named financial metric with an explicit formula. "
            "For growth and CAGR, value_a is current and value_b is previous. "
            "For margin/returns/ratios, value_a is the numerator and value_b is "
            "the denominator."
        ),
        "strict": True,
        "parameters": {
            "type": "object",
            "properties": {
                "metric": {
                    "type": "string",
                    "enum": [
                        "growth_rate",
                        "net_margin",
                        "return_on_assets",
                        "return_on_equity",
                        "current_ratio",
                        "debt_to_equity",
                        "cagr",
                    ],
                },
                "value_a": {"type": "number"},
                "value_b": {"type": "number"},
                "periods": {"type": ["integer", "null"]},
            },
            "required": ["metric", "value_a", "value_b", "periods"],
            "additionalProperties": False,
        },
    },
}

_BINARY_OPERATORS = {
    ast.Add: operator.add,
    ast.Sub: operator.sub,
    ast.Mult: operator.mul,
    ast.Div: operator.truediv,
    ast.FloorDiv: operator.floordiv,
    ast.Mod: operator.mod,
    ast.Pow: operator.pow,
}
_UNARY_OPERATORS = {ast.UAdd: operator.pos, ast.USub: operator.neg}


def split_document(text: str, size: int = 900, overlap: int = 140) -> list[str]:
    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not normalized:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(normalized):
        end = min(start + size, len(normalized))
        if end < len(normalized):
            boundary = normalized.rfind("\n", start, end)
            if boundary <= start + size // 2:
                boundary = normalized.rfind("。", start, end)
            if boundary > start + size // 2:
                end = boundary + 1
        chunks.append(normalized[start:end].strip())
        if end >= len(normalized):
            break
        start = max(end - overlap, start + 1)
    return [chunk for chunk in chunks if chunk]


def extract_document(filename: str, raw: bytes) -> tuple[str, list[dict[str, Any]]]:
    extension = Path(filename).suffix.lower()
    chunks: list[dict[str, Any]] = []

    if extension == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(raw))
        if reader.is_encrypted:
            raise ValueError("Encrypted PDF files are not supported.")
        pages: list[str] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            pages.append(f"[Page {page_number}]\n{text}")
            chunks.extend(
                {"content": chunk, "page": page_number}
                for chunk in split_document(text)
            )
        content = "\n\n".join(pages)
    elif extension == ".docx":
        from docx import Document

        document = Document(BytesIO(raw))
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                blocks.append("\t".join(cell.text.strip() for cell in row.cells))
        content = "\n".join(block for block in blocks if block)
        chunks = [
            {"content": chunk, "page": None}
            for chunk in split_document(content)
        ]
    else:
        try:
            content = raw.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise ValueError("Text documents must use UTF-8 encoding.") from exc
        chunks = [
            {"content": chunk, "page": None}
            for chunk in split_document(content)
        ]

    if not chunks:
        raise ValueError(
            "No readable text was found. Scanned PDFs require OCR before upload."
        )
    if len(content) > MAX_EXTRACTED_CHARACTERS or len(chunks) > MAX_DOCUMENT_CHUNKS:
        raise ValueError(
            "Extracted document text is too large. Split the file into smaller documents."
        )
    return content, chunks


def create_embeddings(texts: list[str]) -> list[list[float]]:
    if not texts:
        return []
    response = client.embeddings.create(
        model=settings.embedding_model,
        input=texts,
    )
    return [item.embedding for item in sorted(response.data, key=lambda item: item.index)]


def retrieve_context(
    user_id: int,
    query: str,
    workspace: str = DEFAULT_WORKSPACE,
    limit: int = 5,
    min_similarity: float = RAG_MIN_SIMILARITY,
) -> list[dict[str, Any]]:
    workspace = validate_workspace(workspace)
    if not database.list_documents(user_id, workspace):
        return []
    query_embedding = create_embeddings([query])[0]
    candidates = database.search_chunks(
        user_id,
        query_embedding,
        workspace=workspace,
        limit=limit,
    )
    return [item for item in candidates if item["score"] >= min_similarity]


def run_space(
    system_prompt: str,
    message: str,
    max_output_tokens: int = 600,
    mode: SpaceRunMode = "lean",
) -> tuple[str, dict[str, int]]:
    """Run a custom user-created space with a hard output token ceiling."""
    response = client.chat.completions.create(
        model=settings.ai_model,
        messages=[
            {
                "role": "system",
                "content": compile_space_system_prompt(system_prompt, mode),
            },
            {"role": "user", "content": message},
        ],
        max_completion_tokens=max(128, min(max_output_tokens, 1_200)),
        store=False,
    )
    usage = response.usage
    input_tokens = int(getattr(usage, "prompt_tokens", 0) or 0)
    output_tokens = int(getattr(usage, "completion_tokens", 0) or 0)
    total_tokens = int(getattr(usage, "total_tokens", input_tokens + output_tokens) or 0)
    return response.choices[0].message.content or "", {
        "input_tokens": input_tokens,
        "output_tokens": output_tokens,
        "total_tokens": total_tokens,
    }


CROSS_EXAM_SCHEMA = {
    "type": "object",
    "properties": {
        "headline": {"type": "string"},
        "executive_summary": {"type": "string"},
        "collisions": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "title": {"type": "string"},
                    "severity": {
                        "type": "string",
                        "enum": ["critical", "high", "medium", "low"],
                    },
                    "confidence": {"type": "integer", "minimum": 0, "maximum": 100},
                    "legal_mechanism": {"type": "string"},
                    "financial_consequence": {"type": "string"},
                    "why_it_matters": {"type": "string"},
                    "legal_source_ids": {
                        "type": "array",
                        "items": {"type": "string"},
                    },
                    "finance_source_ids": {
                        "type": "array",
                        "items": {"type": "string"},
                    },
                    "missing_evidence": {"type": "string"},
                    "next_action": {"type": "string"},
                },
                "required": [
                    "title",
                    "severity",
                    "confidence",
                    "legal_mechanism",
                    "financial_consequence",
                    "why_it_matters",
                    "legal_source_ids",
                    "finance_source_ids",
                    "missing_evidence",
                    "next_action",
                ],
                "additionalProperties": False,
            },
        },
        "stress_scenarios": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "name": {"type": "string"},
                    "trigger": {"type": "string"},
                    "impact_chain": {"type": "string"},
                    "early_warning": {"type": "string"},
                    "response": {"type": "string"},
                },
                "required": [
                    "name",
                    "trigger",
                    "impact_chain",
                    "early_warning",
                    "response",
                ],
                "additionalProperties": False,
            },
        },
        "blind_spots": {
            "type": "array",
            "items": {"type": "string"},
        },
    },
    "required": [
        "headline",
        "executive_summary",
        "collisions",
        "stress_scenarios",
        "blind_spots",
    ],
    "additionalProperties": False,
}


def _numbered_context(
    context: list[dict[str, Any]],
    prefix: str,
) -> str:
    return "\n\n".join(
        (
            f"[{prefix}{index}] {item['name']}"
            + (f", page {item['page']}" if item.get("page") else "")
            + f"\n{item['content']}"
        )
        for index, item in enumerate(context, start=1)
    )


def run_cross_exam(
    focus: str,
    legal_context: list[dict[str, Any]],
    finance_context: list[dict[str, Any]],
) -> dict[str, Any]:
    """Connect legal mechanisms to financial consequences using locked source IDs."""
    instructions = """
You are the FrostFire Cross-Examination Engine. Analyze only the supplied private
contract/compliance excerpts and financial excerpts. Your job is not to summarize
the two folders independently. Build causal bridges: identify a legal mechanism,
show the plausible financial consequence, then state what evidence is missing.

Rules:
- Every collision must cite at least one valid L source ID and one valid F source ID.
- Never invent a clause, number, jurisdiction, market price, or live fact.
- Treat confidence as evidence coverage, not certainty or professional advice.
- Prefer concrete mechanisms such as pricing, payment timing, renewal, termination,
  service levels, indemnity, warranties, covenants, disclosure, or concentration.
- Produce 2-5 collision cards when the evidence permits.
- Produce exactly three stress scenarios: Base, Downside, and Breakpoint. Scenarios
  are reasoned counterfactuals, not forecasts; do not invent numeric impacts.
- Write concise Chinese output.
""".strip()
    prompt = f"""
审查焦点：{focus}

寒冰证据（合同与合规）：
{_numbered_context(legal_context, 'L')}

烈火证据（金融研究）：
{_numbered_context(finance_context, 'F')}
""".strip()
    response = client.chat.completions.create(
        model=settings.ai_model,
        messages=[
            {"role": "system", "content": instructions},
            {"role": "user", "content": prompt},
        ],
        response_format={
            "type": "json_schema",
            "json_schema": {
                "name": "frostfire_cross_exam",
                "strict": True,
                "schema": CROSS_EXAM_SCHEMA,
            },
        },
        max_completion_tokens=1_800,
        store=False,
    )
    content = response.choices[0].message.content or "{}"
    result = json.loads(content)
    if not isinstance(result, dict):
        raise ValueError("Cross-examination result was not a JSON object.")
    return result


def build_messages(
    user_id: int,
    session_id: str,
    context: list[dict[str, Any]],
    workspace: str = DEFAULT_WORKSPACE,
) -> list[dict[str, Any]]:
    workspace = validate_workspace(workspace)
    instructions = [
        "You are a helpful assistant in a persistent chat application.",
        "Answer in Chinese when the user writes Chinese.",
        "Use tools when they materially improve accuracy.",
        "Do not claim a tool was used unless a tool result appears in the conversation.",
        "Clearly separate facts found in sources from analysis and assumptions.",
        "For document-based analysis, prefer the structure: Evidence, Analysis, Gaps, Next checks. Omit a section only when it is genuinely irrelevant.",
    ]
    if workspace == "legal":
        instructions.extend(
            [
                "You are in the Contract and Compliance workspace.",
                "Help review contracts, policies, obligations, deadlines, ambiguities, and compliance evidence.",
                "Cite the supplied source name and page for every material document-based claim when a page is available.",
                "Do not invent clauses, laws, jurisdictions, or regulatory requirements that are absent from the supplied material.",
                "Describe risks as review findings, not definitive legal conclusions, and state when qualified legal review is needed.",
                WORKSPACES["legal"]["boundary"],
            ]
        )
    elif workspace == "finance":
        instructions.extend(
            [
                "You are in the Financial Research workspace.",
                "Help analyze financial reports, announcements, performance, assumptions, and risk disclosures.",
                "Cite the supplied source name and page for every material document-based claim when a page is available.",
                "State currency, units, periods, formulas, and assumptions for calculations.",
                "Never invent current market data or present stale document figures as live data.",
                "Do not provide personalized buy, sell, hold, tax, or portfolio instructions.",
                WORKSPACES["finance"]["boundary"],
            ]
        )
    else:
        instructions.append("You are in the General Document workspace.")
    if context:
        excerpts = "\n\n".join(
            (
                f"[Source: {item['name']}"
                + (f", page {item['page']}" if item.get("page") else "")
                + f"]\n{item['content']}"
            )
            for item in context
        )
        instructions.append(
            "Use the following private knowledge-base excerpts when relevant. "
            "If they do not answer the question, say so rather than inventing facts.\n\n"
            + excerpts
        )

    # Keep the full history in SQLite while sending only the most recent turns
    # back to the model. This gives the regular chat a predictable input-cost
    # ceiling without pretending it has infinite conversational memory.
    stored_messages = database.list_messages(session_id, user_id, limit=16)
    return [
        {"role": "system", "content": "\n\n".join(instructions)},
        *[
            {"role": message["role"], "content": message["content"]}
            for message in stored_messages
        ],
    ]


def _evaluate_node(node: ast.AST) -> float | int:
    if isinstance(node, ast.Expression):
        return _evaluate_node(node.body)
    if isinstance(node, ast.Constant) and isinstance(node.value, (int, float)):
        return node.value
    if isinstance(node, ast.BinOp) and type(node.op) in _BINARY_OPERATORS:
        left = _evaluate_node(node.left)
        right = _evaluate_node(node.right)
        if isinstance(node.op, ast.Pow) and abs(right) > 12:
            raise ValueError("Exponent is too large.")
        result = _BINARY_OPERATORS[type(node.op)](left, right)
        if abs(result) > 1e100:
            raise ValueError("Result is too large.")
        return result
    if isinstance(node, ast.UnaryOp) and type(node.op) in _UNARY_OPERATORS:
        return _UNARY_OPERATORS[type(node.op)](_evaluate_node(node.operand))
    raise ValueError("Unsupported expression.")


def calculate(expression: str) -> dict[str, Any]:
    if len(expression) > 200:
        raise ValueError("Expression is too long.")
    parsed = ast.parse(expression, mode="eval")
    return {"expression": expression, "result": _evaluate_node(parsed)}


def calculate_financial_metric(
    metric: str,
    value_a: float,
    value_b: float,
    periods: int | None = None,
) -> dict[str, Any]:
    if value_b == 0:
        raise ValueError("The denominator or previous value cannot be zero.")

    percent_metrics = {
        "growth_rate": ("(current - previous) / abs(previous)", (value_a - value_b) / abs(value_b)),
        "net_margin": ("net income / revenue", value_a / value_b),
        "return_on_assets": ("net income / total assets", value_a / value_b),
        "return_on_equity": ("net income / total equity", value_a / value_b),
    }
    ratio_metrics = {
        "current_ratio": ("current assets / current liabilities", value_a / value_b),
        "debt_to_equity": ("total debt / total equity", value_a / value_b),
    }

    if metric in percent_metrics:
        formula, result = percent_metrics[metric]
        unit = "%"
        result *= 100
    elif metric in ratio_metrics:
        formula, result = ratio_metrics[metric]
        unit = "x"
    elif metric == "cagr":
        if not periods or periods <= 0:
            raise ValueError("CAGR requires a positive number of periods.")
        if value_a < 0 or value_b <= 0:
            raise ValueError("CAGR requires positive start and non-negative end values.")
        formula = "(current / previous) ** (1 / periods) - 1"
        result = ((value_a / value_b) ** (1 / periods) - 1) * 100
        unit = "%"
    else:
        raise ValueError("Unsupported financial metric.")

    return {
        "metric": metric,
        "formula": formula,
        "value_a": value_a,
        "value_b": value_b,
        "periods": periods,
        "result": round(result, 4),
        "unit": unit,
    }


def get_current_time(timezone_name: str) -> dict[str, str]:
    try:
        zone = ZoneInfo(timezone_name)
    except ZoneInfoNotFoundError as exc:
        raise ValueError("Unknown IANA timezone.") from exc
    now = datetime.now(zone)
    return {
        "timezone": timezone_name,
        "iso": now.isoformat(),
        "display": now.strftime("%Y-%m-%d %H:%M:%S %Z"),
    }


def execute_tool(name: str, arguments: str) -> str:
    try:
        payload = json.loads(arguments or "{}")
        if name == "calculate":
            result = calculate(payload.get("expression", ""))
        elif name == "get_current_time":
            result = get_current_time(payload.get("timezone", "UTC"))
        elif name == "calculate_financial_metric":
            result = calculate_financial_metric(
                payload.get("metric", ""),
                payload.get("value_a"),
                payload.get("value_b"),
                payload.get("periods"),
            )
        else:
            result = {"error": f"Unknown tool: {name}"}
    except (ValueError, SyntaxError, ZeroDivisionError, TypeError) as exc:
        result = {"error": str(exc)}
    return json.dumps(result, ensure_ascii=False)


def tools_for_workspace(workspace: str) -> list[dict[str, Any]]:
    workspace = validate_workspace(workspace)
    return [*BASE_TOOLS, FINANCE_TOOL] if workspace == "finance" else BASE_TOOLS


def run_agent(
    messages: list[dict[str, Any]],
    workspace: str = DEFAULT_WORKSPACE,
) -> tuple[str, list[str], dict[str, int]]:
    working_messages = list(messages)
    tools_used: list[str] = []
    tools = tools_for_workspace(workspace)
    usage_totals = {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0}

    for _ in range(MAX_TOOL_ROUNDS + 1):
        response = client.chat.completions.create(
            model=settings.ai_model,
            messages=working_messages,
            tools=tools,
            tool_choice="auto",
            max_completion_tokens=700,
            store=False,
        )
        response_usage = getattr(response, "usage", None)
        usage_totals["input_tokens"] += int(
            getattr(response_usage, "prompt_tokens", 0) or 0
        )
        usage_totals["output_tokens"] += int(
            getattr(response_usage, "completion_tokens", 0) or 0
        )
        usage_totals["total_tokens"] += int(
            getattr(response_usage, "total_tokens", 0) or 0
        )
        message = response.choices[0].message
        if not message.tool_calls:
            return (
                message.content or "",
                list(dict.fromkeys(tools_used)),
                usage_totals,
            )

        working_messages.append(
            {
                "role": "assistant",
                "content": message.content,
                "tool_calls": [
                    tool_call.model_dump(exclude_none=True)
                    for tool_call in message.tool_calls
                ],
            }
        )
        for tool_call in message.tool_calls:
            name = tool_call.function.name
            tools_used.append(name)
            working_messages.append(
                {
                    "role": "tool",
                    "tool_call_id": tool_call.id,
                    "content": execute_tool(name, tool_call.function.arguments),
                }
            )

    raise RuntimeError("Agent exceeded the tool-call limit.")


def stream_agent(
    messages: list[dict[str, Any]],
    workspace: str = DEFAULT_WORKSPACE,
) -> Generator[dict[str, Any], None, None]:
    working_messages = list(messages)
    tools_used: list[str] = []
    tools = tools_for_workspace(workspace)
    usage_totals = {"input_tokens": 0, "output_tokens": 0, "total_tokens": 0}

    for _ in range(MAX_TOOL_ROUNDS + 1):
        stream = client.chat.completions.create(
            model=settings.ai_model,
            messages=working_messages,
            tools=tools,
            tool_choice="auto",
            max_completion_tokens=700,
            stream=True,
            stream_options={"include_usage": True},
            store=False,
        )
        text_parts: list[str] = []
        tool_buffers: dict[int, dict[str, str]] = {}

        for chunk in stream:
            chunk_usage = getattr(chunk, "usage", None)
            if chunk_usage is not None:
                usage_totals["input_tokens"] += int(
                    getattr(chunk_usage, "prompt_tokens", 0) or 0
                )
                usage_totals["output_tokens"] += int(
                    getattr(chunk_usage, "completion_tokens", 0) or 0
                )
                usage_totals["total_tokens"] += int(
                    getattr(chunk_usage, "total_tokens", 0) or 0
                )
            if not chunk.choices:
                continue
            delta = chunk.choices[0].delta
            if delta.content:
                text_parts.append(delta.content)
                yield {"type": "token", "content": delta.content}
            for partial in delta.tool_calls or []:
                buffer = tool_buffers.setdefault(
                    partial.index,
                    {"id": "", "name": "", "arguments": ""},
                )
                if partial.id:
                    buffer["id"] = partial.id
                if partial.function and partial.function.name:
                    buffer["name"] += partial.function.name
                if partial.function and partial.function.arguments:
                    buffer["arguments"] += partial.function.arguments

        if not tool_buffers:
            yield {
                "type": "done",
                "reply": "".join(text_parts),
                "tools_used": list(dict.fromkeys(tools_used)),
                "usage": usage_totals,
            }
            return

        tool_calls = [
            {
                "id": item["id"],
                "type": "function",
                "function": {
                    "name": item["name"],
                    "arguments": item["arguments"],
                },
            }
            for _, item in sorted(tool_buffers.items())
        ]
        working_messages.append(
            {"role": "assistant", "content": None, "tool_calls": tool_calls}
        )
        for tool_call in tool_calls:
            name = tool_call["function"]["name"]
            tools_used.append(name)
            yield {"type": "tool", "name": name}
            working_messages.append(
                {
                    "role": "tool",
                    "tool_call_id": tool_call["id"],
                    "content": execute_tool(
                        name,
                        tool_call["function"]["arguments"],
                    ),
                }
            )

    raise RuntimeError("Agent exceeded the tool-call limit.")
